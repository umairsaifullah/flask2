from flask import Flask, render_template, request, send_file ,Response
import docx
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from num2words import num2words
import os

app = Flask(__name__)
print('Starting...')


@app.route('/', methods=['GET', 'POST'])
def index():
    print('inside indx...')
    return render_template('index.html')


@app.route('/submit', methods=['GET', 'POST'])
def submit():
    print('inside Submit...')
    example_strings = [request.form['Bankname'], request.form['Bankcode'], request.form['District'],
                       request.form['name'], request.form['cnic'], " ",
                       request.form['case'], request.form['postapplied'], request.form['amount']]

    # Load the existing Word document
    doc = docx.Document('fpscfiletest.docx')

    for o in range(0,4):
        # Get the first table in the document
        table = doc.tables[o]
        table_data = []
        j =1
        for i in range(0, 10):
            if i == 5 or i == 9:
                continue

            if i == 8:
                table.cell(i, j).text = "Rs." + example_strings[i] + "/= (Rupees " + num2words(
                    example_strings[i]) + " Only)"
                paragraph = table.cell(i, j).paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.bold = True
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph_format.line_spacing = Pt(10.35)
                continue
            table.cell(i, j).text = table.cell(i, j).text +" "+ example_strings[i]
            paragraph = table.cell(i, j).paragraphs[0]

            run = paragraph.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.bold = True
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph_format.line_spacing = Pt(10.35)

    # Save the updated document
    doc.save('test_filled.docx')

    print('Converting to PDF...')
   
  
    filename = 'test_filled.docx'
    
    
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(port=5000)
